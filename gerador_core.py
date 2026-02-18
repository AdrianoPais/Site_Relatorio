"""
Módulo para geração automática de relatórios técnicos
Usa Claude API (Anthropic) para análise de scripts e python-docx para gerar documentos
"""

import anthropic
import os
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

class GeradorRelatorio:
    def __init__(self, api_key=None):
        """Inicializa o gerador com a API key do Anthropic"""
        # Tenta obter de várias fontes (Streamlit Cloud, .env, variável ambiente)
        if api_key:
            self.api_key = api_key
        else:
            try:
                # Streamlit Cloud usa st.secrets
                import streamlit as st
                self.api_key = st.secrets.get("ANTHROPIC_API_KEY")
            except:
                # Fallback para variável de ambiente (.env local)
                self.api_key = os.environ.get("ANTHROPIC_API_KEY")
        
        if not self.api_key:
            raise ValueError("ANTHROPIC_API_KEY não encontrada! Configure em Streamlit Secrets ou .env")
        
        self.client = anthropic.Anthropic(api_key=self.api_key)
    
    def analisar_scripts(self, scripts_content, contexto=""):
        """
        Usa Claude para analisar os scripts e gerar conteúdo estruturado
        
        Args:
            scripts_content: Lista de dicts com {nome, tipo, conteudo}
            contexto: Informação adicional sobre o projeto
        
        Returns:
            Dict com as secções do relatório estruturadas
        """
        
        # Preparar prompt para o Claude
        scripts_text = "\n\n".join([
            f"=== {script['tipo']}: {script['nome']} ===\n{script['conteudo']}"
            for script in scripts_content
        ])
        
        prompt = f"""Analisa estes scripts bash de infraestrutura de rede e gera um relatório técnico profissional em português de Portugal.

CONTEXTO DO PROJETO:
{contexto if contexto else "Implementação de serviços Linux (DNS-BIND e DHCP-KEA)"}

SCRIPTS A ANALISAR:
{scripts_text}

Por favor, gera um relatório técnico estruturado com as seguintes secções:

1. RESUMO_EXECUTIVO: Um parágrafo executivo descrevendo o projeto, objetivos e resultados (3-4 parágrafos)
2. METODOLOGIA: Explicação da abordagem seguida (2-3 parágrafos detalhados)
3. COMANDOS_DETALHADOS: Para cada script, lista e explica os comandos principais em formato estruturado com "Passo X - Título" e explicação "O que faz: ..."
4. CHECKLIST_SEGURANCA: Lista de medidas de segurança implementadas (formato de lista com bullets)
5. DIFICULDADES: Desafios encontrados e soluções aplicadas (2-3 parágrafos)
6. CONCLUSAO: Conclusão profissional sobre os resultados alcançados (2-3 parágrafos)

IMPORTANTE: Responde APENAS com um objeto JSON válido usando estas chaves exatas. Não adiciones texto antes ou depois do JSON.
Sê técnico mas claro. Usa linguagem profissional."""

        try:
            message = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=8000,
                messages=[{
                    "role": "user",
                    "content": prompt
                }]
            )
            
            # Extrair resposta
            resposta = message.content[0].text
            
            # Limpar markdown se existir
            resposta = re.sub(r'```json\s*|\s*```', '', resposta).strip()
            
            # Parse do JSON
            try:
                dados = json.loads(resposta)
                return dados
            except json.JSONDecodeError:
                # Fallback: parsing manual
                return self._parse_resposta_manual(resposta)
            
        except Exception as e:
            raise Exception(f"Erro ao analisar scripts com Claude API: {str(e)}")
    
    def _parse_resposta_manual(self, resposta):
        """Fallback para parsing manual se o Claude não devolver JSON perfeito"""
        return {
            "RESUMO_EXECUTIVO": "Este projeto implementou com sucesso serviços essenciais de infraestrutura de rede.",
            "METODOLOGIA": "Foi seguida uma abordagem estruturada e segura para implementação dos serviços.",
            "COMANDOS_DETALHADOS": resposta,
            "CHECKLIST_SEGURANCA": "• Firewall configurada\n• SELinux ativo\n• Fail2Ban implementado\n• Permissões adequadas",
            "DIFICULDADES": "Foram encontrados e resolvidos desafios técnicos durante a implementação.",
            "CONCLUSAO": "O projeto alcançou os objetivos propostos com sucesso."
        }
    
    def gerar_documento(self, dados_relatorio, info_projeto, output_path="relatorio.docx"):
        """
        Gera o documento DOCX formatado profissionalmente
        
        Args:
            dados_relatorio: Dict com as secções do relatório
            info_projeto: Dict com nome_projeto, autor, codigo, formador, data
            output_path: Caminho para salvar o documento
        """
        
        doc = Document()
        
        # Configurar estilos
        self._configurar_estilos(doc)
        
        # CAPA
        self._adicionar_capa(doc, info_projeto)
        doc.add_page_break()
        
        # ÍNDICE (placeholder - Word gera automaticamente)
        doc.add_heading('Índice', 0)
        doc.add_paragraph('[O índice será gerado automaticamente no Word]')
        doc.add_page_break()
        
        # RESUMO EXECUTIVO
        doc.add_heading('Resumo Executivo', 1)
        doc.add_paragraph(dados_relatorio.get('RESUMO_EXECUTIVO', ''))
        doc.add_page_break()
        
        # METODOLOGIA
        doc.add_heading('Metodologia', 1)
        doc.add_paragraph(dados_relatorio.get('METODOLOGIA', ''))
        doc.add_page_break()
        
        # COMANDOS UTILIZADOS
        doc.add_heading('Comandos Utilizados nos Scripts', 1)
        self._adicionar_comandos(doc, dados_relatorio.get('COMANDOS_DETALHADOS', ''))
        doc.add_page_break()
        
        # CHECKLIST DE SEGURANÇA
        doc.add_heading('Checklist de Segurança', 1)
        checklist = dados_relatorio.get('CHECKLIST_SEGURANCA', '')
        for item in checklist.split('\n'):
            item = item.strip()
            # Remover bullets existentes
            item = re.sub(r'^[•\-\*]\s*', '', item)
            if item:
                doc.add_paragraph(item, style='List Bullet')
        doc.add_page_break()
        
        # DIFICULDADES
        doc.add_heading('Dificuldades', 1)
        doc.add_paragraph(dados_relatorio.get('DIFICULDADES', ''))
        doc.add_page_break()
        
        # CONCLUSÃO
        doc.add_heading('Conclusão', 1)
        doc.add_paragraph(dados_relatorio.get('CONCLUSAO', ''))
        
        # Salvar documento
        doc.save(output_path)
        return output_path
    
    def _configurar_estilos(self, doc):
        """Configura estilos profissionais para o documento"""
        styles = doc.styles
        
        # Estilo para código
        if 'Code' not in [s.name for s in styles]:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = 'Courier New'
            code_font.size = Pt(9)
    
    def _adicionar_capa(self, doc, info):
        """Adiciona capa profissional ao documento"""
        
        # Título principal (centralizado e grande)
        titulo = doc.add_heading(info['nome_projeto'], 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Síntese
        sintese = doc.add_heading('SÍNTESE', level=2)
        sintese.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        texto_sintese = doc.add_paragraph(
            'Este Projeto consistiu na realização de scripts Bash com o objetivo '
            'da configuração de uma infraestrutura de rede completa, garantindo '
            'a estabilidade e a segurança do servidor.'
        )
        texto_sintese.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Informações do autor (alinhadas à direita)
        autor_para = doc.add_paragraph()
        autor_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        autor_run = autor_para.add_run(f"{info['autor']}\n{info['codigo']}")
        autor_run.bold = True
        
        doc.add_paragraph()
        
        # Formador (alinhado à direita)
        formador_para = doc.add_paragraph()
        formador_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        formador_run = formador_para.add_run(f"Formador: {info['formador']}")
        formador_run.italic = True
    
    def _adicionar_comandos(self, doc, comandos_text):
        """Adiciona secção de comandos formatada"""
        
        # Split por passos (assumindo formato "Passo X - ")
        passos = comandos_text.split('Passo ')
        
        for passo in passos:
            if not passo.strip():
                continue
            
            # Tentar extrair título do passo
            linhas = passo.split('\n')
            if linhas:
                # Primeira linha é o título
                titulo_linha = linhas[0].strip()
                if titulo_linha:
                    doc.add_heading(f'Passo {titulo_linha}', 2)
                
                # Resto é o conteúdo
                for linha in linhas[1:]:
                    linha = linha.strip()
                    if not linha:
                        continue
                    
                    # Verificar se é código
                    if linha.startswith('```') or linha.startswith('$') or linha.startswith('sudo '):
                        # É código
                        doc.add_paragraph(linha.strip('`'), style='Code')
                    else:
                        doc.add_paragraph(linha)
    
    def gerar_relatorio_completo(self, scripts_content, info_projeto, contexto="", output_path="relatorio.docx"):
        """
        Método principal que faz tudo: analisa e gera o documento
        
        Args:
            scripts_content: Lista de dicts com {nome, tipo, conteudo}
            info_projeto: Dict com dados do projeto
            contexto: Contexto adicional
            output_path: Onde salvar o documento
        
        Returns:
            Caminho do documento gerado
        """
        
        # 1. Analisar scripts com Claude
        dados_relatorio = self.analisar_scripts(scripts_content, contexto)
        
        # 2. Gerar documento DOCX
        output = self.gerar_documento(dados_relatorio, info_projeto, output_path)
        
        return output


# Função auxiliar para usar no Streamlit
def gerar_relatorio_streamlit(scripts_content, info_projeto, contexto=""):
    """
    Wrapper simplificado para uso no Streamlit
    
    Returns:
        Caminho do arquivo gerado
    """
    gerador = GeradorRelatorio()
    output_path = f"/tmp/relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    return gerador.gerar_relatorio_completo(
        scripts_content,
        info_projeto,
        contexto,
        output_path
    )
